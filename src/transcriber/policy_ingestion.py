"""Utilities for ingesting policy documents and templates."""

from __future__ import annotations

import io
import json
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional

from openai import OpenAI

from .utils import stable_filename


@dataclass
class PolicyDocument:
    name: str
    text: str
    source_bytes: bytes
    metadata: Dict[str, Any]


@dataclass
class PolicyTemplate:
    name: str
    schema: Dict[str, Any]
    created_at: str
    source_policy_name: Optional[str] = None


class DocumentExtractor:
    """Extract text from supported document formats."""

    @staticmethod
    def load(uploaded_file) -> PolicyDocument:
        name = uploaded_file.name
        raw = uploaded_file.read()
        uploaded_file.seek(0)

        suffix = name.lower().rsplit(".", 1)[-1] if "." in name else ""

        if suffix in {"txt", "md", "json"}:
            text = DocumentExtractor._decode_text(raw)
        elif suffix == "pdf":
            text = DocumentExtractor._extract_pdf(io.BytesIO(raw))
        elif suffix in {"docx"}:
            text = DocumentExtractor._extract_docx(io.BytesIO(raw))
        else:
            text = DocumentExtractor._decode_text(raw)

        return PolicyDocument(
            name=name,
            text=text.strip(),
            source_bytes=raw,
            metadata={"extension": suffix},
        )

    @staticmethod
    def _decode_text(data: bytes) -> str:
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="ignore")

    @staticmethod
    def _extract_pdf(buffer: io.BytesIO) -> str:
        try:
            from pypdf import PdfReader
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("pypdf is required to process PDF files") from exc

        reader = PdfReader(buffer)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    @staticmethod
    def _extract_docx(buffer: io.BytesIO) -> str:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required to process DOCX files") from exc

        document = docx.Document(buffer)
        parts = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(parts)


class TemplateBuilder:
    """Create structured templates for compliance scoring."""

    def __init__(self, client: OpenAI, model: str):
        self.client = client
        self.model = model

    def build_from_policy(
        self,
        policy: PolicyDocument,
        instructions: Optional[str] = None,
        scorecard_structure: Optional[Dict[str, Any]] = None,
    ) -> PolicyTemplate:
        prompt = self._compose_prompt(policy.text, instructions, scorecard_structure)
        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "system",
                    "content": "You convert fire department radio policies into structured scoring templates."
                },
                {"role": "user", "content": prompt},
            ],
            response_format={"type": "json_schema", "json_schema": self._schema_contract()},
        )

        payload = self._extract_response_json(response)
        return PolicyTemplate(
            name=payload.get("template_name", policy.name),
            schema=payload,
            created_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            source_policy_name=policy.name,
        )

    @staticmethod
    def load_template_file(uploaded_file) -> PolicyTemplate:
        name = uploaded_file.name
        raw = uploaded_file.read()
        uploaded_file.seek(0)
        data = json.loads(raw.decode("utf-8"))
        return PolicyTemplate(
            name=data.get("template_name", name),
            schema=data,
            created_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        )

    @staticmethod
    def _compose_prompt(policy_text: str, instructions: Optional[str], scorecard: Optional[Dict[str, Any]]) -> str:
        base = (
            "Given the following fire department radio policy, produce a JSON template "
            "with scoring categories, criteria, pass/fail guidance, and weighting."
        )
        if instructions:
            base += f"\nAdditional user instructions: {instructions}\n"
        if scorecard:
            base += "\nIncorporate elements of the user's existing scorecard table (columns, weights, criteria) provided below."
            base += "\nExisting scorecard JSON:\n" + json.dumps(scorecard)[:8000]
        return base + "\nPolicy text:\n" + policy_text[:10000]

    @staticmethod
    def _schema_contract() -> Dict[str, Any]:
        return {
            "name": "fire_policy_template",
            "schema": {
                "type": "object",
                "properties": {
                    "template_name": {"type": "string"},
                    "overall_weighting": {"type": "object"},
                    "categories": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "weight": {"type": "number"},
                                "criteria": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "id": {"type": "string"},
                                            "description": {"type": "string"},
                                            "guidance": {"type": "string"},
                                            "weight": {"type": "number"},
                                        },
                                        "required": ["id", "description"],
                                    },
                                },
                            },
                            "required": ["name", "criteria"],
                        },
                    },
                },
                "required": ["categories"],
            },
        }

    @staticmethod
    def _extract_response_json(response: Any) -> Dict[str, Any]:
        try:
            output = response.output_text
        except AttributeError as exc:  # pragma: no cover
            raise RuntimeError("Unexpected response structure from template generation") from exc
        return json.loads(output)


class TemplateStore:
    """Persist and retrieve policy templates from disk."""

    def __init__(self, root: Path = Path("transcripts/policies")):
        self.root = root
        self.root.mkdir(parents=True, exist_ok=True)

    def save(self, template: PolicyTemplate) -> Path:
        base = Path(template.name).stem or "template"
        fname = stable_filename(base.replace(" ", "_"), "template", "json")
        dst = self.root / fname
        dst.write_text(json.dumps(template.schema, indent=2), encoding="utf-8")
        return dst

    def list_templates(self) -> Dict[str, Dict[str, Any]]:
        templates: Dict[str, Dict[str, Any]] = {}
        for path in self.root.glob("*.json"):
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                continue
            name = data.get("template_name") or path.stem
            templates[name] = data
        return templates
